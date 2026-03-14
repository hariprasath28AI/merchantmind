"""Generate MerchantMind demo cheat-sheet as a DOCX file."""
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn


def set_cell_shading(cell, color_hex):
    """Set background color of a table cell."""
    shading = cell._element.get_or_add_tcPr()
    sh = shading.makeelement(qn("w:shd"), {
        qn("w:fill"): color_hex,
        qn("w:val"): "clear",
    })
    shading.append(sh)


def add_heading_styled(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    return h


def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def build():
    doc = Document()

    # --- Page margins ---
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ─── TITLE ────────────────────────────────────────────
    title = doc.add_heading("MerchantMind", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.size = Pt(32)
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Autonomous Reconciliation Agent for Pine Labs")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    run.italic = True

    tagline = doc.add_paragraph()
    tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = tagline.add_run("Pine Labs Playground Hackathon  ·  March 14, 2026")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_paragraph()  # spacer

    # ─── ONE-LINER ────────────────────────────────────────
    add_heading_styled(doc, "The Problem", level=1)
    p = doc.add_paragraph()
    run = p.add_run(
        "Merchants processing thousands of daily transactions across UPI, cards, BNPL, and EMIs "
        "face fragmented, manual reconciliation. Settlement delays (up to 5 days), separate chargeback "
        "cycles, and multi-instrument complexity mean finance teams spend hours daily finding mismatches "
        "— and still miss anomalies causing direct revenue loss."
    )
    run.font.size = Pt(11)

    p = doc.add_paragraph()
    run = p.add_run("Pine Labs' own blog calls this process \"tedious\" and \"error-prone.\" "
                     "Their solution today: a dashboard. Our solution: an agent that acts.")
    run.bold = True
    run.font.size = Pt(11)

    # ─── SOLUTION ─────────────────────────────────────────
    add_heading_styled(doc, "The Solution: MerchantMind", level=1)
    doc.add_paragraph(
        "An autonomous merchant intelligence agent that listens, detects, reasons, acts, and reports — "
        "using the full Pine Labs API stack."
    )

    steps = [
        ("1. Listens ", "to Pine Labs webhook events in real-time (payment.captured, refund.created, etc.)"),
        ("2. Pulls ", "settlement data via REST API and cross-references transactions"),
        ("3. Detects ", "anomalies using 10 rule-based patterns (shortfalls, duplicates, fraud velocity, etc.)"),
        ("4. Reasons ", "about each anomaly using Claude AI — explains WHY it's a problem and WHAT to do"),
        ("5. Acts ", "autonomously — auto-refunds below ₹500, blocks duplicates, raises fraud alerts"),
        ("6. Reports ", "every decision transparently: \"Found 10 issues. Fixed 2 automatically. 4 flagged. 4 need your review.\""),
    ]
    for bold, rest in steps:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(bold)
        r.bold = True
        p.add_run(rest)

    # ─── PINE LABS API USAGE ──────────────────────────────
    add_heading_styled(doc, "Full Pine Labs API Integration (All 3 Layers)", level=1)

    table = doc.add_table(rows=4, cols=3)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ["Layer", "API", "How MerchantMind Uses It"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)

    rows_data = [
        ("Webhooks", "payment.captured\npayment.failed\nrefund.created\norder.cancelled",
         "Real-time trigger — agent wakes on event, broadcasts to dashboard, starts scan"),
        ("Settlement REST API", "GET /api/settlements/v1/list\nGET by UTR\nPOST /api/auth/v1/token",
         "Reconciliation core — pulls settlement data, cross-references against transactions to find shortfalls/excess"),
        ("MCP Server Tools", "create_refund\ncancel_order\nget_order_by_order_id\nsearch_transaction",
         "Action layer — executes auto-refunds, cancels duplicates, verifies post-action. Real API calls with real sandbox order IDs"),
    ]
    for i, (layer, api, usage) in enumerate(rows_data):
        table.rows[i + 1].cells[0].text = layer
        table.rows[i + 1].cells[1].text = api
        table.rows[i + 1].cells[2].text = usage
        for j in range(3):
            for p in table.rows[i + 1].cells[j].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)

    doc.add_paragraph()  # spacer

    # ─── ANOMALY DETECTION ────────────────────────────────
    add_heading_styled(doc, "10 Anomaly Detection Patterns", level=1)

    anomaly_table = doc.add_table(rows=11, cols=4)
    anomaly_table.style = "Light Grid Accent 1"
    anomaly_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ["#", "Anomaly", "Severity", "Auto-Action"]
    for i, h in enumerate(headers):
        cell = anomaly_table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)

    anomalies = [
        ("1", "Settlement Shortfall", "MEDIUM/HIGH", "Auto-refund if ≤ ₹500"),
        ("2", "Duplicate Refund", "HIGH", "Block + flag for review"),
        ("3", "Refund Velocity Fraud", "CRITICAL", "Fraud alert, pause actions"),
        ("4", "Over-Settlement", "MEDIUM", "Flag for review"),
        ("5", "Late Settlement (>3 days)", "HIGH", "Escalate to support"),
        ("6", "Refund Exceeds Capture", "HIGH", "Block refund"),
        ("7", "Duplicate Order", "HIGH", "Cancel duplicate"),
        ("8", "High-Value Outlier (>10x avg)", "MEDIUM", "Hold for confirmation"),
        ("9", "Partial Capture Mismatch", "MEDIUM", "Flag over-settlement"),
        ("10", "Midnight Burst (>5 txns 12-5AM)", "MEDIUM", "Flag unusual activity"),
    ]
    for i, (num, name, sev, action) in enumerate(anomalies):
        row = anomaly_table.rows[i + 1]
        row.cells[0].text = num
        row.cells[1].text = name
        row.cells[2].text = sev
        row.cells[3].text = action
        for j in range(4):
            for p in row.cells[j].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)

    doc.add_paragraph()

    # ─── COMPETITIVE ANALYSIS ─────────────────────────────
    add_heading_styled(doc, "What Pine Labs Has vs. What We Add", level=1)

    comp_table = doc.add_table(rows=8, cols=3)
    comp_table.style = "Light Grid Accent 1"
    comp_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(["Capability", "Pine Labs Today", "MerchantMind"]):
        cell = comp_table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)

    comp_rows = [
        ("Anomaly Detection", "Basic discrepancy matching", "10 pattern-specific rules + Claude AI reasoning"),
        ("Fraud Detection", "Gateway-level (millisecond auth)", "Merchant-level: velocity, midnight burst, duplicate refund"),
        ("Autonomous Action", "None — dashboard only", "Auto-refund, auto-cancel, block & flag"),
        ("Reasoning Transparency", "None", "Per-anomaly Claude reasoning with full audit trail"),
        ("Real-time Streaming", "Transaction status updates", "Full agent pipeline streamed live via WebSocket"),
        ("Merchant Control", "Manual review everything", "Approve/dismiss with threshold-gated auto-actions"),
        ("Webhook-Driven", "Events available but no consumer", "Agent wakes on events, scans, notifies dashboard"),
    ]
    for i, (cap, today, mm) in enumerate(comp_rows):
        row = comp_table.rows[i + 1]
        row.cells[0].text = cap
        row.cells[1].text = today
        row.cells[2].text = mm
        for j in range(3):
            for p in row.cells[j].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)

    doc.add_paragraph()

    # ─── WHY THIS WINS ───────────────────────────────────
    add_heading_styled(doc, "Why This Wins for Pine Labs", level=1)

    value_props = [
        ("We build what their OpenAI partnership promises — on their own stack. ",
         "The Feb 2026 OpenAI deal targets internal settlement ops. MerchantMind is merchant-facing "
         "autonomous reconciliation. Pine Labs can ship this to every merchant without depending on OpenAI's timeline."),
        ("First real consumer of all 3 API layers. ",
         "No one else has demonstrated Webhooks + Settlement REST + MCP Server working together in a single "
         "autonomous agent loop. We prove their API strategy works end-to-end."),
        ("We solve the problem they blog about but haven't solved. ",
         "Pine Labs blogs say manual reconciliation is \"tedious\" and \"error-prone.\" "
         "They offer a dashboard. We offer an agent that detects, reasons, and acts."),
        ("\"AI will govern payment stacks\" — we're the proof. ",
         "At the India AI Impact Summit 2026, Pine Labs said AI will move from sitting alongside payment stacks "
         "to governing them. MerchantMind is that governance layer."),
        ("Merchant stickiness multiplier. ",
         "A merchant using Pine Labs + MerchantMind gets autonomous money protection. That's a switching cost "
         "no competitor can match. Pine Labs evolves from payment processor to merchant intelligence platform."),
    ]
    for bold, rest in value_props:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(bold)
        r.bold = True
        p.add_run(rest)

    # ─── TECH STACK ───────────────────────────────────────
    add_heading_styled(doc, "Tech Stack", level=1)

    stack_table = doc.add_table(rows=7, cols=2)
    stack_table.style = "Light Grid Accent 1"
    stack_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    stack_headers = ["Layer", "Technology"]
    for i, h in enumerate(stack_headers):
        cell = stack_table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)

    stack_rows = [
        ("Agent Orchestration", "LangGraph (StateGraph with conditional edges)"),
        ("LLM Reasoning", "Claude Sonnet 4 via Anthropic API"),
        ("Backend", "Python + FastAPI + WebSocket streaming"),
        ("Pine Labs Integration", "REST API + MCP Server + Webhooks (all 3 layers)"),
        ("Frontend Dashboard", "React with real-time WebSocket updates"),
        ("Data", "500 mock transactions + 10 planted anomalies + 2 real sandbox orders"),
    ]
    for i, (layer, tech) in enumerate(stack_rows):
        stack_table.rows[i + 1].cells[0].text = layer
        stack_table.rows[i + 1].cells[1].text = tech
        for j in range(2):
            for p in stack_table.rows[i + 1].cells[j].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)

    doc.add_paragraph()

    # ─── DEMO SCRIPT ──────────────────────────────────────
    add_heading_styled(doc, "Demo Script (5 Minutes)", level=1)

    demo_steps = [
        ("Minute 1 — The Problem (Hook)",
         "\"A D2C merchant on Pine Labs processes 2,000 transactions a day. Their finance team spends "
         "3 hours every morning reconciling settlements manually. They still miss things — and it costs "
         "them money. Today we're going to show you what happens when that problem disappears.\""),
        ("Minute 2 — Live Agent Scan",
         "Trigger scan. Show real-time streaming: agent loads 500 transactions, detects 10 anomalies, "
         "reasons about each one with Claude AI. Point out the live reasoning appearing on each card — "
         "\"Claude AI\" badge means real LLM analysis, not just rules."),
        ("Minute 3 — Real Pine Labs API Call",
         "Expand the settlement shortfall card. Show the real Pine Labs order ID. "
         "\"This isn't a mock — the agent called Pine Labs' actual refund API on a real sandbox order.\" "
         "Show the API call badge in the expanded card."),
        ("Minute 4 — Fraud Detection + Merchant Control",
         "Show the CRITICAL fraud velocity alert. \"5 refunds to the same card in 60 minutes — ₹2,500 total. "
         "Agent caught it and paused everything. Nothing moves without merchant approval.\" "
         "Click Approve on one card. Click Dismiss on another. Show cross-tab sync."),
        ("Minute 5 — Business Value Close",
         "\"We used all three Pine Labs API layers: Webhooks as triggers, Settlement REST for data, "
         "MCP for execution. Pine Labs said at the AI Summit: AI will govern payment stacks. "
         "MerchantMind is that governance layer — deployable to every Pine Labs merchant tomorrow.\""),
    ]
    for title, script in demo_steps:
        p = doc.add_paragraph()
        r = p.add_run(title)
        r.bold = True
        r.font.size = Pt(11)
        r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
        p2 = doc.add_paragraph(script)
        p2.paragraph_format.left_indent = Cm(1)
        for run in p2.runs:
            run.font.size = Pt(10)
            run.italic = True

    # ─── JUDGE Q&A ────────────────────────────────────────
    add_heading_styled(doc, "Anticipated Judge Questions", level=1)

    qas = [
        ("Q: What if the LLM makes a wrong decision and issues a wrong refund?",
         "All auto-actions have a configurable threshold — ₹500 by default. Above that, every action requires "
         "merchant approval. The agent acts as an advisor, not a dictator. Plus, every decision includes "
         "Claude's reasoning and confidence score for full audit trail."),
        ("Q: How does this scale to thousands of merchants?",
         "Each merchant gets an isolated agent instance. LangGraph state is stateless per invocation — "
         "horizontally scalable on AWS Lambda or ECS. The LLM layer scales via Anthropic API with no infra overhead."),
        ("Q: Why not just use rule-based automation?",
         "Rules catch what you anticipate. The LLM catches what you don't. When a new fraud pattern "
         "emerges — like refund velocity to rotating cards — the agent reasons about it contextually, "
         "not from a hardcoded rule. Our 10 rules are the fast filter; Claude is the intelligent second opinion."),
        ("Q: How is this different from Pine Labs' OpenAI partnership?",
         "Their OpenAI deal (Feb 2026) targets internal settlement operations — cutting clearing time from "
         "hours to minutes for Pine Labs' own ops team. MerchantMind is merchant-facing: every merchant gets "
         "an autonomous agent protecting their money. Complementary, not competitive."),
        ("Q: Did you actually use the Pine Labs API?",
         "Yes. We created real orders on the Pine Labs sandbox (MID 121478), and the agent makes real API "
         "calls to the refund and cancel endpoints. We use all 3 API layers: Webhooks for triggers, "
         "Settlement REST for data, and MCP-style calls for execution."),
    ]
    for q, a in qas:
        p = doc.add_paragraph()
        r = p.add_run(q)
        r.bold = True
        r.font.size = Pt(10)
        p2 = doc.add_paragraph(a)
        p2.paragraph_format.left_indent = Cm(1)
        for run in p2.runs:
            run.font.size = Pt(10)

    # ─── SOURCES ──────────────────────────────────────────
    add_heading_styled(doc, "Pine Labs Sources Referenced", level=1)

    sources = [
        "Pine Labs Blog: Payment Reconciliation vs Settlement — pinelabs.com/blog/understanding-the-difference-between-payment-gateway-settlement-and-reconciliation",
        "Pine Labs Blog: Best Practices for Payment Reconciliation — pinelabs.com/blog/your-guide-to-the-best-practices-of-payment-reconciliation",
        "Pine Labs Blog: Retailers Guide to Reconciliation — pinelabs.com/blog/retailers-guide-to-payment-reconciliation-process",
        "Plural: Streamline Payment Settlements for SMBs — pluralonline.com/streamline-payment-settlements-for-smbs-best-practices/",
        "Pine Labs MCP Server Docs — developer.pinelabsonline.com/docs/mcp-server",
        "Plural Blog: MCP Server — pluralonline.com/pine-labs-mcp-server-making-payment-integration-easy-for-everyone/",
        "TechCrunch: OpenAI-Pine Labs Partnership (Feb 2026) — techcrunch.com/2026/02/18/openai-deepens-india-push-with-pine-labs-fintech-partnership/",
        "Pine Labs Blog: AI & Payments at India AI Impact Summit 2026 — pinelabs.com/blog/blog-india-ai-impact-summit-2026-payment-infrastructure",
    ]
    for s in sources:
        p = doc.add_paragraph(s, style="List Bullet")
        for run in p.runs:
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # ─── SAVE ─────────────────────────────────────────────
    out_path = "MerchantMind_Demo_CheatSheet.docx"
    doc.save(out_path)
    print(f"Saved: {out_path}")
    return out_path


if __name__ == "__main__":
    build()
